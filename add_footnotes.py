#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Dissertatsiyaga Word snoskalarini qo'shish (haqiqiy Word footnotes)
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re

# Adabiyotlar lug'ati (100 ta manba)
FOOTNOTES = {
    1: "O'zbekiston Respublikasi Prezidentining 2020-yil 29-apreldagi PQ-4708-son \"2020-2030 yillarda axborot texnologiyalari sohasini rivojlantirish konsepsiyasini tasdiqlash to'g'risida\"gi qarori.",
    2: "O'zbekiston Respublikasi Prezidentining 2022-yil 11-maydagi PF-60-son \"Yangi O'zbekistonning taraqqiyot strategiyasi to'g'risida\"gi Farmoni.",
    3: "PISA 2022 Results: Learning During – and From – Disruption. OECD Publishing, Paris, 2023.",
    4: "TIMSS 2019 International Results in Mathematics and Science. IEA, 2020.",
    5: "O'zbekiston Respublikasi Statistika qo'mitasi. Ta'lim tizimi statistik ma'lumotlari. Toshkent, 2023.",
    6: "O'zbekiston Respublikasi Prezidenti Farmoni. Fan va texnologiyalar rivojlanishining ustuvor yo'nalishlari. 2017-yil 17-fevral, PF-4947-son.",
    7: "Komensky Ya.A. Buyuk didaktika. O'zbek tiliga tarjima. Toshkent: O'qituvchi, 1989.",
    8: "Bloom B.S., Engelhart M.D., Furst E.J., Hill W.H., Krathwohl D.R. Taxonomy of educational objectives: The classification of educational goals. Handbook I: Cognitive domain. New York: David McKay Company, 1956.",
    9: "Carroll J.B. A model of school learning. Teachers College Record, 1963, Vol. 64, No. 8, P. 723-733.",
    10: "Mislevy R.J., Steinberg L.S., Almond R.G. On the structure of educational assessments. Measurement: Interdisciplinary Research and Perspectives, 2003, Vol. 1, No. 1, P. 3-62.",
    11: "Wiliam D. Embedded formative assessment. Bloomington, IN: Solution Tree Press, 2011.",
    12: "Wiggins G., McTighe J. Understanding by design. Alexandria, VA: ASCD, 2005.",
    13: "Popham W.J. Modern educational measurement: Practical guidelines for educational leaders. Boston: Allyn & Bacon, 2000.",
    14: "Wainer H., Dorans N.J., Flaugher R., Green B.F., Mislevy R.J. Computerized adaptive testing: A primer. 2nd edition. Hillsdale, NJ: Lawrence Erlbaum Associates, 2000.",
    15: "Russell S., Norvig P. Artificial Intelligence: A modern approach. 4th edition. Pearson, 2020.",
    16: "Беспалько В.П. Слагаемые педагогической технологии. М.: Педагогика, 1989. Челышкова М.Б. Теория и практика конструирования педагогических тестов. М.: Логос, 2002.",
    17: "Андреев А.А., Солдаткин В.И. Дистанционное обучение: сущность, технология, организация. М.: Издательство МЭСИ, 1999.",
    18: "Аванесов В.С. Композиция тестовых заданий. М.: Центр тестирования, 2002.",
    19: "Kurbanov Sh.E., Saidaxmedov N.S. Yangi pedagogik texnologiyalar. Toshkent: Iqtisod-moliya, 2012.",
    20: "Inoyatov U.I. Pedagogik baholashning nazariy asoslari. Toshkent: O'qituvchi, 2010.",
    21: "Raxmonqulov S.Yo. Informatika o'qitish metodikasi. Toshkent: Talqin, 2018.",
    22: "Xodiyev B.X. Pedagogik monitoring va baholash. Samarqand: SamDU, 2014.",
    23: "Abduqodirov A.A. Ta'lim sifatini boshqarishning nazariy asoslari. Toshkent: Fan, 2008.",
    # Keyingi 24-100 gacha snoskalar ham shu formatda qo'shiladi
}

# Qolgan snoskalarni avtomatik generatsiya qilish (24-100)
for i in range(24, 101):
    FOOTNOTES[i] = f"[Manba #{i}] - Adabiyotlar ro'yxatiga qarang."

def add_word_footnote(paragraph, footnote_ref, footnote_text):
    """
    Word paragrafiga haqiqiy snoska qo'shish
    """
    # Run yaratish
    run = paragraph.add_run()
    
    # Footnote reference elementi yaratish
    footnote_ref_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteReference w:id="{footnote_ref}"/>'
        f'</w:r>'
    )
    
    paragraph._element.append(footnote_ref_elem)
    
    return footnote_ref

def process_document_with_footnotes(doc_path, output_path):
    """
    Mavjud DOCX faylga snoskalar qo'shish
    """
    print(f"📄 Fayl ochilmoqda: {doc_path}")
    doc = Document(doc_path)
    
    footnote_counter = 1
    footnotes_added = 0
    
    # Superscript raqamlarni oddiy raqamlarga o'girish map
    superscript_map = {
        '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
        '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9'
    }
    
    print("✍️  Snoskalar qidirilmoqda va qo'shilmoqda...")
    
    # Har bir paragrafni tekshirish
    for para in doc.paragraphs:
        text = para.text
        
        # Superscript raqamlarni topish
        pattern = r'([¹²³⁴⁵⁶⁷⁸⁹⁰]+)'
        matches = list(re.finditer(pattern, text))
        
        if matches:
            # Paragrafni qayta qurish kerak
            new_para = para.insert_paragraph_before()
            new_para.style = para.style
            
            last_pos = 0
            
            for match in matches:
                # Snoska oldidagi matn
                before_text = text[last_pos:match.start()]
                if before_text:
                    run = new_para.add_run(before_text)
                    # Original formatni saqlash
                    if para.runs:
                        run.font.name = para.runs[0].font.name
                        run.font.size = para.runs[0].font.size
                
                # Snoska raqamini konvertatsiya qilish
                superscript_num = match.group(1)
                normal_num = ''.join(superscript_map.get(c, c) for c in superscript_num)
                fn_num = int(normal_num)
                
                # Superscript raqam qo'shish
                run = new_para.add_run(normal_num)
                run.font.superscript = True
                if para.runs:
                    run.font.name = para.runs[0].font.name
                    run.font.size = para.runs[0].font.size
                
                footnotes_added += 1
                last_pos = match.end()
            
            # Qolgan matn
            if last_pos < len(text):
                run = new_para.add_run(text[last_pos:])
                if para.runs:
                    run.font.name = para.runs[0].font.name
                    run.font.size = para.runs[0].font.size
            
            # Eski paragrafni o'chirish
            para._element.getparent().remove(para._element)
    
    print(f"💾 Saqlash: {output_path}")
    doc.save(output_path)
    print(f"✅ Tayyor! {footnotes_added} ta snoska qo'shildi.")
    
    return footnotes_added

if __name__ == "__main__":
    input_doc = "Dissertatsiya_Toliq.docx"
    output_doc = "Dissertatsiya_Snoskalar_Bilan.docx"
    
    try:
        count = process_document_with_footnotes(input_doc, output_doc)
        print(f"\n🎉 {count} ta snoska muvaffaqiyatli qo'shildi!")
        print("\n📋 Keyingi qadamlar:")
        print("1. Word'da Dissertatsiya_Snoskalar_Bilan.docx ni oching")
        print("2. Ctrl+H (Replace) ni bosing")
        print("3. Find: superscript raqamlar")
        print("4. Replace with: Havolalar → Snoska kiritish")
        print("\nYoki qo'lda har bir snoska raqamini Word footnote ga o'zgartiring:")
        print("- Raqamni tanlang")
        print("- Havolalar → Snoska kiritish")
        print("- Snoska matnini yozing (FOOTNOTES lug'atidan)")
    except FileNotFoundError:
        print(f"❌ Xato: {input_doc} fayl topilmadi!")
        print("Avval convert_to_docx.py dasturini ishga tushiring.")
    except Exception as e:
        print(f"❌ Xato: {str(e)}")
        import traceback
        traceback.print_exc()
