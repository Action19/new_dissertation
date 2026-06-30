#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Dissertatsiyani Markdown formatdan DOCX formatga o'tkazish dasturi
OAK talablariga muvofiq
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styles(doc):
    """OAK talablariga mos uslublar yaratish"""
    
    # Normal (asosiy matn) uslubi
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(14)
    
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    # BOB sarlavhasi uslubi
    try:
        style_heading1 = doc.styles.add_style('BOB_Sarlavha', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_heading1 = doc.styles['BOB_Sarlavha']
    
    font_h1 = style_heading1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(14)
    font_h1.bold = True
    
    paragraph_format_h1 = style_heading1.paragraph_format
    paragraph_format_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_h1.space_before = Pt(12)
    paragraph_format_h1.space_after = Pt(12)
    paragraph_format_h1.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Bo'lim sarlavhasi uslubi
    try:
        style_heading2 = doc.styles.add_style('Bolim_Sarlavha', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_heading2 = doc.styles['Bolim_Sarlavha']
    
    font_h2 = style_heading2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(14)
    font_h2.bold = True
    
    paragraph_format_h2 = style_heading2.paragraph_format
    paragraph_format_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format_h2.space_before = Pt(6)
    paragraph_format_h2.space_after = Pt(6)
    paragraph_format_h2.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format_h2.first_line_indent = Cm(0)

def setup_page(doc):
    """Sahifa parametrlarini sozlash (OAK talablari)"""
    section = doc.sections[0]
    
    # A4 o'lchami
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    
    # Chegaralar: Yuqori/Quyi 20mm, Chap 30mm, O'ng 15mm
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

def process_line(doc, line, current_style='Normal'):
    """Har bir qatorni qayta ishlash va to'g'ri formatda qo'shish"""
    
    line = line.strip()
    
    if not line:
        doc.add_paragraph()
        return current_style
    
    # # bilan boshlanadigan sarlavhalar
    if line.startswith('# '):
        text = line[2:].strip()
        # BOB yoki KIRISH, XULOSA kabi asosiy bo'limlar
        if any(keyword in text.upper() for keyword in ['BOB', 'KIRISH', 'XULOSA', 'MUNDARIJA', 'ADABIYOTLAR', 'ILOVALAR']):
            p = doc.add_paragraph(text, style='BOB_Sarlavha')
        else:
            p = doc.add_paragraph(text, style='Bolim_Sarlavha')
        return 'Heading'
    
    elif line.startswith('## '):
        text = line[3:].strip()
        p = doc.add_paragraph(text, style='Bolim_Sarlavha')
        return 'Heading'
    
    elif line.startswith('### '):
        text = line[4:].strip()
        p = doc.add_paragraph(text)
        p.runs[0].bold = True
        p.paragraph_format.first_line_indent = Cm(0)
        return 'Heading'
    
    # Ro'yxatlar (bullet points)
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        return 'List'
    
    # Raqamlangan ro'yxatlar
    elif re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line)
        p = doc.add_paragraph(text, style='List Number')
        p.paragraph_format.first_line_indent = Cm(0)
        return 'List'
    
    # Jadval chegaralari (Unicode box-drawing)
    elif line.startswith('┌') or line.startswith('├') or line.startswith('└') or line.startswith('│'):
        # Jadvallarni oddiy matn sifatida qo'shish
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        return 'Table'
    
    # Kod bloklari
    elif line.startswith('```'):
        return 'Code'
    
    # Oddiy matn
    else:
        # Qalin matn (**text** yoki __text__)
        if '**' in line or '__' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|__.*?__)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('__') and part.endswith('__'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph(line)
        
        return 'Normal'

def convert_markdown_to_docx(input_file, output_file):
    """Asosiy konvertatsiya funksiyasi"""
    
    print(f"📄 Fayl o'qilmoqda: {input_file}")
    
    # Markdown faylni o'qish
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Yangi Word hujjati yaratish
    doc = Document()
    
    # Sahifa parametrlarini sozlash
    setup_page(doc)
    
    # Uslublarni yaratish
    create_styles(doc)
    
    print("✍️  DOCX formatga konvertatsiya qilinmoqda...")
    
    # Qatorma-qator qayta ishlash
    lines = content.split('\n')
    current_style = 'Normal'
    in_code_block = False
    
    for i, line in enumerate(lines):
        # Kod bloklarini boshqarish
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
            continue
        
        # Oddiy qatorlarni qayta ishlash
        current_style = process_line(doc, line, current_style)
        
        # Progress ko'rsatish (har 100 qatorda)
        if (i + 1) % 100 == 0:
            print(f"   {i + 1}/{len(lines)} qator qayta ishlandi...")
    
    # Faylni saqlash
    print(f"💾 Saqlash: {output_file}")
    doc.save(output_file)
    print(f"✅ Tayyor! {output_file} yaratildi.")
    print(f"📊 Jami {len(lines)} qator qayta ishlandi.")

if __name__ == "__main__":
    input_file = "Dissertatsiya_Toliq.md"
    output_file = "Dissertatsiya_Toliq.docx"
    
    try:
        convert_markdown_to_docx(input_file, output_file)
        print("\n🎉 Dissertatsiya muvaffaqiyatli DOCX formatga o'tkazildi!")
        print("\n📋 Keyingi qadamlar:")
        print("1. Dissertatsiya_Toliq.docx faylni oching")
        print("2. Mundarija va sahifa raqamlarini qo'shing")
        print("3. Jadvallar va rasmlarni tekshiring")
        print("4. Snoskalarni formatlang")
        print("5. PDF formatda saqlang")
    except FileNotFoundError:
        print(f"❌ Xato: {input_file} fayl topilmadi!")
    except Exception as e:
        print(f"❌ Xato yuz berdi: {str(e)}")
