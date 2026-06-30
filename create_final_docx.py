#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
To'liq formatlangan dissertatsiya yaratish
Barcha matnlar, snoskalar, jadvallar bilan
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class DissertationConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.footnote_num = 1
        
    def setup_document(self):
        """OAK talablariga muvofiq"""
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
    
    def add_para(self, text, **kwargs):
        """Paragraf qo'shish"""
        bold = kwargs.get('bold', False)
        italic = kwargs.get('italic', False)
        font_size = kwargs.get('font_size', 14)
        alignment = kwargs.get('alignment', 'justify')
        indent = kwargs.get('indent', True)
        space_before = kwargs.get('space_before', 0)
        space_after = kwargs.get('space_after', 0)
        
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        # Shrift
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        
        # Cyrillic
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Format
        fmt = para.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        if indent and alignment == 'justify':
            fmt.first_line_indent = Cm(1.25)
        
        if alignment == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'justify':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif alignment == 'left':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if space_before > 0:
            fmt.space_before = Pt(space_before)
        if space_after > 0:
            fmt.space_after = Pt(space_after)
        
        return para


    
    def process_line(self, line):
        """Har bir qatorni qayta ishlash"""
        line = line.strip()
        
        if not line or line == '---':
            return None
        
        # Sarlavhalar
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            if text and text != '---':
                return ('h1', text)
        
        elif line.startswith('## '):
            text = line[3:].strip()
            return ('h2', text)
        
        elif line.startswith('### '):
            text = line[4:].strip()
            return ('h3', text)
        
        elif line.startswith('#### '):
            text = line[5:].strip()
            return ('h4', text)
        
        # Bold matn
        elif line.startswith('**') and line.endswith('**'):
            text = line.strip('*')
            return ('bold', text)
        
        # Jadval qatori
        elif line.startswith('|'):
            return ('table', line)
        
        # Ro'yxat
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            return ('list', text)
        
        # Raqamlangan ro'yxat
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            return ('numbered', text)
        
        # Kod bloki boshlashi
        elif line.startswith('```'):
            return ('code_start', '')
        
        # Snoska ta'rifi
        elif line.startswith('[^'):
            return None  # O'tkazib yuboramiz
        
        # Oddiy matn
        else:
            # Snoskalarni qayta ishlash
            text = self.process_footnotes(line)
            # Markdown ni tozalash
            text = self.clean_markdown(text)
            if len(text) > 3:
                return ('text', text)
        
        return None
    
    def process_footnotes(self, text):
        """Snoskalarni superscript raqamlarga o'zgartirish"""
        # [^1] -> ¹
        def get_superscript(num):
            superscripts = '⁰¹²³⁴⁵⁶⁷⁸⁹'
            if num < 10:
                return superscripts[num]
            else:
                return ''.join(superscripts[int(d)] for d in str(num))
        
        text = re.sub(r'\[\^(\d+)\]', lambda m: get_superscript(int(m.group(1))), text)
        return text
    
    def clean_markdown(self, text):
        """Markdown belgilarini tozalash"""
        # Bold
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Italic
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Havolalar
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        return text
    
    def convert(self, md_file):
        """Markdown faylni DOCX ga o'tkazish"""
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        in_code_block = False
        code_lines = []
        
        for line in lines:
            if in_code_block:
                if line.strip().startswith('```'):
                    # Kod bloki tugadi
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        self.add_para(code_text, font_size=11, alignment='left', indent=False)
                    in_code_block = False
                    code_lines = []
                else:
                    code_lines.append(line.rstrip())
                continue
            
            result = self.process_line(line)
            
            if result is None:
                continue
            
            typ, content = result
            
            if typ == 'h1':
                self.add_para(content, bold=True, font_size=16, alignment='center',
                            indent=False, space_before=12, space_after=12)
            
            elif typ == 'h2':
                self.add_para(content, bold=True, font_size=14, alignment='left',
                            indent=False, space_before=10, space_after=10)
            
            elif typ == 'h3':
                self.add_para(content, bold=True, font_size=14, alignment='left',
                            indent=False, space_before=8, space_after=8)
            
            elif typ == 'h4':
                self.add_para(content, bold=True, font_size=13, alignment='left',
                            indent=False, space_before=6, space_after=6)
            
            elif typ == 'bold':
                self.add_para(content, bold=True, alignment='center', indent=False)
            
            elif typ == 'table':
                self.add_para(content, font_size=12, alignment='left', indent=False)
            
            elif typ == 'list':
                self.add_para('• ' + content, alignment='left', indent=False)
            
            elif typ == 'numbered':
                self.add_para(content, alignment='left')
            
            elif typ == 'code_start':
                in_code_block = True
                code_lines = []
            
            elif typ == 'text':
                self.add_para(content)
        
        # Sahifa raqamlash
        self.add_page_numbers()
    
    def add_page_numbers(self):
        """Sahifa raqamlash"""
        section = self.doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    def save(self, output_file):
        """Saqlash"""
        self.doc.save(output_file)
        para_count = len([p for p in self.doc.paragraphs if p.text.strip()])
        print(f"✅ Yaratildi: {output_file}")
        print(f"📝 Paragraflar: {para_count}")
        print(f"📄 Taxminiy sahifalar: {para_count // 30 + 5}")

if __name__ == "__main__":
    print("🚀 To'liq dissertatsiya yaratish boshlandi...")
    
    converter = DissertationConverter()
    converter.convert("Dissertatsiya_Yangi.md")
    converter.save("Dissertatsiya_Yangi_Final.docx")
    
    print("✨ Muvaffaqiyatli yakunlandi!")
