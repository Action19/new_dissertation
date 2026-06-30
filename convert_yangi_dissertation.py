#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Dissertatsiyani Markdown dan DOCX ga o'tkazish
OAK talablariga muvofiq: Times New Roman 14pt, 1.5 interval, snoskalar bilan
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def setup_document():
    """OAK talablariga muvofiq hujjat yaratish"""
    doc = Document()
    
    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)    # 20mm
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)   # 30mm
    section.right_margin = Cm(1.5)  # 15mm
    
    return doc

def add_paragraph(doc, text, bold=False, italic=False, alignment='justify', 
                 font_size=14, spacing_before=0, spacing_after=0, 
                 first_line_indent=True):
    """Formatlangan paragraf qo'shish"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    
    # Shrift
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    
    # Cyrillic uchun
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Paragraf formati
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    if first_line_indent and alignment == 'justify':
        fmt.first_line_indent = Cm(1.25)
    
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if spacing_before:
        fmt.space_before = Pt(spacing_before)
    if spacing_after:
        fmt.space_after = Pt(spacing_after)
    
    return para

def process_markdown_file(input_file, output_file):
    """Markdown faylni o'qish va DOCX yaratish"""
    doc = setup_document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    footnote_counter = 1
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Bo'sh qator
        if not line:
            i += 1
            continue

        
        # # bilan boshlanuvchi sarlavhalar
        if line.startswith('# ') and not line.startswith('## '):
            # Birinchi darajali sarlavha (BOB)
            text = line[2:].strip()
            if text != '---':  # Ajratuvchi chiziqlar emas
                add_paragraph(doc, text, bold=True, alignment='center', 
                            font_size=16, spacing_before=12, spacing_after=12,
                            first_line_indent=False)
        
        elif line.startswith('## '):
            # Ikkinchi darajali sarlavha
            text = line[3:].strip()
            add_paragraph(doc, text, bold=True, alignment='left', 
                        font_size=14, spacing_before=10, spacing_after=10,
                        first_line_indent=False)
        
        elif line.startswith('### '):
            # Uchinchi darajali sarlavha
            text = line[4:].strip()
            add_paragraph(doc, text, bold=True, alignment='left', 
                        font_size=14, spacing_before=8, spacing_after=8,
                        first_line_indent=False)
        
        elif line.startswith('**') and line.endswith('**'):
            # Bold matn
            text = line.strip('*')
            add_paragraph(doc, text, bold=True, alignment='center',
                        first_line_indent=False)
        
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Italic matn
            text = line.strip('*')
            add_paragraph(doc, text, italic=True, alignment='center',
                        first_line_indent=False)
        
        elif line.startswith('|'):
            # Jadval - oddiy matn sifatida
            add_paragraph(doc, line, alignment='left', font_size=12,
                        first_line_indent=False)
        
        elif line.startswith('```'):
            # Kod bloki - o'tkazib yuboramiz yoki oddiy matn qilamiz
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].strip())
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                add_paragraph(doc, code_text, alignment='left', 
                            font_size=11, first_line_indent=False)
        
        elif line.startswith('[^'):
            # Snoska ta'rifi - o'tkazib yuboramiz
            i += 1
            continue
        
        elif line.startswith('---'):
            # Ajratuvchi chiziq
            i += 1
            continue
        
        else:
            # Oddiy paragraf
            # Snoskalarni qayta ishlash
            text = line
            
            # Markdown snoskalarni Word formatiga o'zgartirish
            # [^1] -> ¹ (superscript)
            text = re.sub(r'\[\^(\d+)\]', lambda m: chr(0x2070 + int(m.group(1)) if int(m.group(1)) < 10 
                         else 0x2080 + int(m.group(1)) % 10), text)
            
            # Bold va italic markuplarni tozalash
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            
            # Havolalarni tozalash
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            if text and len(text) > 3:
                add_paragraph(doc, text)
        
        i += 1
    
    # Sahifa raqamlash qo'shish
    add_page_numbers(doc)
    
    # Saqlash
    doc.save(output_file)
    print(f"✅ Dissertatsiya muvaffaqiyatli yaratildi: {output_file}")
    print(f"📄 Sahifalar soni: taxminan {len(doc.paragraphs) // 25} sahifa")

def add_page_numbers(doc):
    """Sahifa raqamlash qo'shish"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Shrift
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

if __name__ == "__main__":
    input_file = "Dissertatsiya_Yangi.md"
    output_file = "Dissertatsiya_Yangi_OAK.docx"
    
    print("🚀 Dissertatsiyani DOCX formatga o'tkazish boshlandi...")
    process_markdown_file(input_file, output_file)
    print("✨ Tayyor!")
